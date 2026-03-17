import csv
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import unquote

# ----------------- DOI parsing -----------------
DOI_URL_RE = re.compile(r"https?://(?:dx\.)?doi\.org/([^\s<>\]\)\"']+)", re.IGNORECASE)
DOI_CORE_RE = re.compile(r"(10\.\d{4,9}/[^\s<>\]\)\"']+)", re.IGNORECASE)


def extract_doi_core(s: str) -> str:
    if not s:
        return ""
    s2 = s.strip()
    m = DOI_URL_RE.search(s2)
    if m:
        return m.group(1)
    m = DOI_CORE_RE.search(s2)
    if m:
        return m.group(1)
    return ""


def normalize_doi(s: str) -> str:
    core = extract_doi_core(s)
    if not core:
        core = (s or "").strip()
    core = core.strip()
    core = core.replace("doi:", "").strip()
    core = core.rstrip(".,;:)]}>\"'")
    core = unquote(core)
    core2 = extract_doi_core(core)
    if core2:
        core = core2
    return core.lower()


def extract_dois_from_text(txt: str) -> list[str]:
    found = []
    for m in DOI_URL_RE.finditer(txt or ""):
        found.append(normalize_doi(m.group(1)))
    for m in DOI_CORE_RE.finditer(txt or ""):
        found.append(normalize_doi(m.group(1)))
    out, seen = [], set()
    for d in found:
        if d and d not in seen:
            seen.add(d)
            out.append(d)
    return out


# ----------------- XML-safe text cleaning -----------------
# XML 1.0 allows tab, LF, CR, and most printable Unicode, but rejects
# control chars like \x00-\x08, \x0B, \x0C, \x0E-\x1F.
XML_ILLEGAL_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")


def xml_safe(s: str) -> str:
    if not s:
        return ""
    return XML_ILLEGAL_RE.sub("", s)


# ----------------- PDF text extraction (best-effort) -----------------
def extract_text_from_pdf(pdf_path: Path) -> str:
    text_parts = []

    try:
        import pdfplumber  # type: ignore
        with pdfplumber.open(str(pdf_path)) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t:
                    text_parts.append(t)
        return xml_safe("\n\n".join(text_parts))
    except Exception:
        pass

    try:
        from PyPDF2 import PdfReader  # type: ignore
        reader = PdfReader(str(pdf_path))
        for page in reader.pages:
            t = page.extract_text() or ""
            if t:
                text_parts.append(t)
        return xml_safe("\n\n".join(text_parts))
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from {pdf_path.name}: {e}") from e


# ----------------- CSV mapping (Citation col1, DOI col2) -----------------
def load_citation_csv(csv_path: Path) -> tuple[dict[str, str], dict[str, str]]:
    doi_to_citation: dict[str, str] = {}
    doi_to_url: dict[str, str] = {}

    with csv_path.open("r", encoding="utf-8-sig", newline="") as f:
        rows = list(csv.reader(f))

    if not rows:
        return doi_to_citation, doi_to_url

    first = [c.strip().lower() for c in rows[0][:2]]
    start_idx = 1 if ("citation" in first[0] or "doi" in first[1]) else 0

    for r in rows[start_idx:]:
        if len(r) < 2:
            continue
        citation = xml_safe((r[0] or "").strip())
        doi_raw = (r[1] or "").strip()
        doi = normalize_doi(doi_raw)
        if not doi:
            continue
        doi_to_citation[doi] = citation
        doi_to_url[doi] = f"https://doi.org/{doi}"

    return doi_to_citation, doi_to_url


# ----------------- PDF filename normalization + matching -----------------
def normalize_filename_for_doi_search(name: str) -> str:
    s = unquote(name).lower()
    s = s.replace("https=", "https:")
    s = s.replace("http=", "http:")
    s = s.replace("]", "/")
    s = s.replace("\\", "/")
    s = re.sub(r"/{2,}", "/", s)
    return s


def index_pdfs_by_doi(folder: Path) -> dict[str, Path]:
    doi_to_pdf: dict[str, Path] = {}
    for p in folder.glob("*.pdf"):
        norm = normalize_filename_for_doi_search(p.stem)
        m = DOI_CORE_RE.search(norm)
        if m:
            doi = normalize_doi(m.group(1))
            doi_to_pdf.setdefault(doi, p)
            continue
        if "doi.org/" in norm:
            after = norm.split("doi.org/", 1)[1]
            m2 = DOI_CORE_RE.search(after)
            if m2:
                doi = normalize_doi(m2.group(1))
                doi_to_pdf.setdefault(doi, p)
    return doi_to_pdf


def match_pdf_for_doi(folder: Path, doi: str, doi_index: dict[str, Path]) -> Path | None:
    if doi in doi_index:
        return doi_index[doi]

    variants = {
        doi,
        doi.replace("/", "_"),
        doi.replace("/", "-"),
        doi.replace("/", "]"),
        "doi.org/" + doi,
        "doi.org]" + doi.replace("/", "]"),
    }

    for p in folder.glob("*.pdf"):
        fname_norm = normalize_filename_for_doi_search(p.name)
        fname_bracket = fname_norm.replace("/", "]")
        if any(v in fname_norm for v in variants) or any(v in fname_bracket for v in variants):
            return p
    return None


# ----------------- Keyword matching (standalone-token style) -----------------
def compile_keyword_regex(keywords: list[str]) -> re.Pattern | None:
    kws = [k.strip() for k in keywords if k and k.strip()]
    if not kws:
        return None
    kws = sorted(kws, key=len, reverse=True)
    escaped = [re.escape(k) for k in kws]
    pat = r"(?<![0-9A-Za-z])(" + "|".join(escaped) + r")(?![0-9A-Za-z])"
    return re.compile(pat, re.IGNORECASE)


@dataclass(frozen=True)
class Span:
    start: int
    end: int


def merge_spans(spans: list[Span]) -> list[Span]:
    if not spans:
        return []
    spans_sorted = sorted(spans, key=lambda s: s.start)
    merged = [spans_sorted[0]]
    for s in spans_sorted[1:]:
        last = merged[-1]
        if s.start <= last.end:
            merged[-1] = Span(last.start, max(last.end, s.end))
        else:
            merged.append(s)
    return merged


def find_snippet_spans(text: str, kw_re: re.Pattern, window: int = 400) -> list[Span]:
    spans: list[Span] = []
    for m in kw_re.finditer(text):
        s = max(0, m.start() - window)
        e = min(len(text), m.end() + window)
        spans.append(Span(s, e))
    return merge_spans(spans)


# ----------------- Cleaning: unwrap lines + remove in-text citations -----------------
CIT_REMOVE_PATTERNS = [
    re.compile(r"\([^\)]*(?:et\s+al\.?|19\d{2}|20\d{2})[^\)]*\)", re.IGNORECASE),
    re.compile(r"\[[^\]]*(?:et\s+al\.?|19\d{2}|20\d{2})[^\]]*\]", re.IGNORECASE),
    re.compile(r"\{[^\}]*(?:et\s+al\.?|19\d{2}|20\d{2})[^\}]*\}", re.IGNORECASE),
]


def remove_intext_citations(s: str) -> str:
    out = s
    for _ in range(3):
        before = out
        for pat in CIT_REMOVE_PATTERNS:
            out = pat.sub("", out)
        if out == before:
            break
    return out


def clean_snippet(snippet: str) -> str:
    s = xml_safe(snippet or "")
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = s.replace("-\n", "-")
    s = re.sub(r"[\n\t]+", " ", s)
    s = remove_intext_citations(s)
    s = re.sub(r"[ ]{2,}", " ", s).strip()
    s = re.sub(r"\s+([,.;:])", r"\1", s)
    s = xml_safe(s)
    return s


# ----------------- Stop at last "references" substring -----------------
def truncate_at_last_references(text: str) -> str:
    """
    Keep only content before the last occurrence of substring 'references' (case-insensitive).
    If not found, return text unchanged.
    """
    if not text:
        return text
    lower = text.lower()
    idx = lower.rfind("references")
    if idx == -1:
        return text
    return text[:idx]


# ----------------- Word export -----------------
def write_docx(output_path: Path, results: list[dict], kw_re: re.Pattern) -> None:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.enum.dml import MSO_THEME_COLOR  # type: ignore
    from docx.shared import Pt, RGBColor  # type: ignore

    doc = Document()
    doc.add_heading("Keyword Snippets from DOI PDFs", level=0)

    for item in results:
        citation = xml_safe(item["citation"] or "(Missing citation in CSV)")
        doi_url = xml_safe(item["doi_url"])
        snippets = [xml_safe(s) for s in item["snippets"]]

        cit_p = doc.add_paragraph()
        cit_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cit_run = cit_p.add_run(citation)
        cit_run.font.size = Pt(12)
        cit_run.font.color.theme_color = MSO_THEME_COLOR.ACCENT_1

        doi_p = doc.add_paragraph()
        doi_run = doi_p.add_run(doi_url)
        doi_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        for snip in snippets:
            snip = xml_safe(snip)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            pos = 0
            for m in kw_re.finditer(snip):
                if m.start() > pos:
                    p.add_run(xml_safe(snip[pos:m.start()]))
                r = p.add_run(xml_safe(snip[m.start():m.end()]))
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW
                pos = m.end()

            if pos < len(snip):
                p.add_run(xml_safe(snip[pos:]))

        doc.add_paragraph("")

    doc.save(str(output_path))


# ----------------- Main -----------------
def main():
    folder = Path(".").resolve()
    txt_path = folder / "Citation_Export.txt"
    csv_path = folder / "doi_citations.csv"

    if not txt_path.exists():
        print(f"ERROR: {txt_path.name} not found in {folder}")
        sys.exit(1)
    if not csv_path.exists():
        print(f"ERROR: {csv_path.name} not found in {folder}")
        sys.exit(1)

    keywords = ["L1B", "Level-1B", "ECO_L1B_GEO", "L1B_GEO", "Level-1B Geolocation product"]
    kw_re = compile_keyword_regex(keywords)
    if kw_re is None:
        print("ERROR: No keywords configured.")
        sys.exit(1)

    txt = txt_path.read_text(encoding="utf-8", errors="ignore")
    txt = xml_safe(txt)

    dois_in_txt = extract_dois_from_text(txt)
    if not dois_in_txt:
        print("No DOIs found in Citation_Export.txt.")
        sys.exit(0)

    doi_to_citation, doi_to_url = load_citation_csv(csv_path)
    doi_index = index_pdfs_by_doi(folder)

    results = []
    missing_pdfs = []
    missing_csv = []

    for doi_raw in dois_in_txt:
        doi = normalize_doi(doi_raw)

        pdf_path = match_pdf_for_doi(folder, doi, doi_index)
        if pdf_path is None:
            missing_pdfs.append(doi)
            continue

        citation = doi_to_citation.get(doi, "")
        doi_url = doi_to_url.get(doi, f"https://doi.org/{doi}")
        if doi not in doi_to_url:
            missing_csv.append(doi)

        try:
            text = extract_text_from_pdf(pdf_path)
            text = xml_safe(text)
        except Exception as e:
            print(f"WARNING: {e}")
            text = ""

        if not text:
            snippets = []
        else:
            text_pre_refs = truncate_at_last_references(text)
            spans = find_snippet_spans(text_pre_refs, kw_re, window=400)
            snippets = [clean_snippet(text_pre_refs[s.start:s.end]) for s in spans]
            snippets = [s for s in snippets if s.strip()]

        print(f"Processed: {pdf_path.name} | snippets: {len(snippets)}")

        if not snippets:
            continue

        results.append({
            "doi_url": doi_url,
            "citation": citation,
            "snippets": snippets,
        })

    if not results:
        print("\nNo matching snippets found in any processed PDF.")
        sys.exit(0)

    out_docx = folder / "DOI_Keyword_Extracts.docx"
    write_docx(out_docx, results, kw_re)
    print(f"\nWrote: {out_docx}")

    if missing_pdfs:
        print("\nWARNING: Could not match these DOIs to any PDF filename:")
        for d in missing_pdfs:
            print(f"  - {d}")

    if missing_csv:
        print("\nNOTE: These DOIs were processed but not found in doi_citations.csv:")
        for d in sorted(set(missing_csv)):
            print(f"  - {d}")


if __name__ == "__main__":
    main()
